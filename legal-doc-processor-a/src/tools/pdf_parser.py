import re
from typing import Dict, Any
from pypdf import PdfReader
from docx import Document
import logging

logger = logging.getLogger(__name__)

class DocumentParser:
    """MCP Tool 1: Parse various document formats"""
    
    @staticmethod
    def parse_pdf(file_path: str) -> Dict[str, Any]:
        """Parse PDF document and extract text with metadata"""
        try:
            reader = PdfReader(file_path)
            
            # Extract metadata
            metadata = {
                'page_count': len(reader.pages),
                'author': reader.metadata.get('/Author', 'Unknown'),
                'title': reader.metadata.get('/Title', 'Unknown'),
                'creation_date': str(reader.metadata.get('/CreationDate', 'Unknown'))
            }
            
            # Extract text with page information
            full_text = []
            page_texts = {}
            
            for page_num, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text()
                page_texts[f"page_{page_num}"] = page_text
                full_text.append(f"[Page {page_num}]\n{page_text}\n")
            
            result = {
                'full_text': '\n'.join(full_text),
                'metadata': metadata,
                'page_texts': page_texts,
                'format': 'pdf'
            }
            
            logger.info(f"Successfully parsed PDF: {file_path} ({metadata['page_count']} pages)")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def parse_docx(file_path: str) -> Dict[str, Any]:
        """Parse DOCX document"""
        try:
            doc = Document(file_path)
            
            # Extract text with paragraph structure
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            metadata = {
                'page_count': len(doc.sections),
                'paragraph_count': len(doc.paragraphs),
                'title': doc.core_properties.title or 'Unknown',
                'author': doc.core_properties.author or 'Unknown'
            }
            
            result = {
                'full_text': '\n\n'.join(full_text),
                'metadata': metadata,
                'format': 'docx'
            }
            
            logger.info(f"Successfully parsed DOCX: {file_path}")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def parse_document(file_path: str) -> Dict[str, Any]:
        """Auto-detect format and parse document"""
        if file_path.lower().endswith('.pdf'):
            return DocumentParser.parse_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            return DocumentParser.parse_docx(file_path)
        elif file_path.lower().endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return {
                'full_text': text,
                'metadata': {'format': 'txt'},
                'format': 'txt'
            }
        else:
            raise ValueError(f"Unsupported file format: {file_path}")
    
    @staticmethod
    def estimate_complexity(text: str) -> float:
        """Estimate document complexity for model routing"""
        # Factors: length, legal jargon density, sentence complexity
        
        words = text.split()
        word_count = len(words)
        
        # Legal terms indicator
        legal_terms = [
            'whereas', 'hereinafter', 'notwithstanding', 'indemnify',
            'liability', 'jurisdiction', 'arbitration', 'covenant',
            'warranties', 'representations', 'severability', 'governing law'
        ]
        
        legal_term_count = sum(1 for word in words if word.lower() in legal_terms)
        legal_density = legal_term_count / max(word_count, 1)
        
        # Sentence complexity (average sentence length)
        sentences = re.split(r'[.!?]+', text)
        avg_sentence_length = word_count / max(len(sentences), 1)
        
        # Calculate complexity score (0-1)
        length_score = min(word_count / 10000, 1.0) * 0.3
        legal_score = min(legal_density * 10, 1.0) * 0.4
        complexity_score = min(avg_sentence_length / 50, 1.0) * 0.3
        
        total_complexity = length_score + legal_score + complexity_score
        
        return min(total_complexity, 1.0)